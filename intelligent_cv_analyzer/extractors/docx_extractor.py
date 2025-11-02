"""
DOCX Text Extractor
Phase 3: Core Application Implementation

This module handles text extraction from DOCX files using python-docx library.
It provides robust error handling and text cleaning for CV analysis.
"""

import os
import time
import logging
from typing import Dict, Any, List

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DOCXExtractor:
    """
    DOCX text extraction class with error handling and text cleaning.
    """
    
    def __init__(self):
        """Initialize the DOCX extractor."""
        # Note: python-docx only supports .docx (Office 2007+), not legacy .doc files
        self.supported_extensions = ['.docx']
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        start_time = time.time()
        
        result = {
            'success': False,
            'text': '',
            'metadata': {},
            'error': None,
            'extraction_time': 0.0
        }
        
        try:
            # Check for legacy .doc files (not supported by python-docx)
            if file_path.lower().endswith('.doc') and not file_path.lower().endswith('.docx'):
                result['error'] = "Legacy .doc format not supported. Only .docx files (Office 2007+) are supported. Please convert to .docx or use alternative extraction."
                result['extraction_time'] = time.time() - start_time
                return result
            
            # Validate file
            if not self._validate_file(file_path):
                result['error'] = "Invalid DOCX file"
                return result
            
            # Check if DOCX library is available
            if not DOCX_AVAILABLE:
                result['error'] = "DOCX extraction library not available"
                return result
            
            # Extract text using python-docx
            text = self._extract_text_from_docx(file_path)
            
            # Clean and normalize text
            cleaned_text = self._clean_text(text)
            
            # Get file metadata
            metadata = self._get_file_metadata(file_path)
            
            result.update({
                'success': True,
                'text': cleaned_text,
                'metadata': metadata,
                'extraction_time': time.time() - start_time
            })
            
        except Exception as e:
            result['error'] = f"DOCX extraction failed: {str(e)}"
            result['extraction_time'] = time.time() - start_time
        
        return result
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        # Join all text parts
        full_text = ' '.join(text_parts)
        
        return full_text
    
    def _validate_file(self, file_path: str) -> bool:
        """
        Validate DOCX file before extraction.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            True if file is valid, False otherwise
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return False
            
            # Check file extension
            if not any(file_path.lower().endswith(ext) for ext in self.supported_extensions):
                return False
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                return False
            
            # Check if file is readable
            if not os.access(file_path, os.R_OK):
                return False
            
            return True
            
        except Exception:
            return False
    
    def _clean_text(self, raw_text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            raw_text: Raw text extracted from DOCX
            
        Returns:
            Cleaned and normalized text
        """
        if not raw_text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(raw_text.split())
        
        # Remove special characters that might interfere with matching
        # Keep alphanumeric, spaces, and common punctuation
        import re
        text = re.sub(r'[^\w\s\-.,;:!?()]', ' ', text)
        
        # Normalize multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing file metadata
        """
        metadata = {
            'filename': os.path.basename(file_path),
            'file_size': 0,
            'file_path': file_path,
            'extraction_timestamp': time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        try:
            metadata['file_size'] = os.path.getsize(file_path)
            
            # Try to extract document properties
            if DOCX_AVAILABLE:
                try:
                    doc = Document(file_path)
                    core_props = doc.core_properties
                    
                    if core_props.title:
                        metadata['title'] = core_props.title
                    if core_props.author:
                        metadata['author'] = core_props.author
                    if core_props.created:
                        metadata['created'] = core_props.created.isoformat()
                    if core_props.modified:
                        metadata['modified'] = core_props.modified.isoformat()
                        
                except Exception as e:
                    logging.getLogger(__name__).warning("Failed to extract DOCX core properties: %s", e, exc_info=True)
                    
        except Exception as e:
            logging.getLogger(__name__).warning("Failed to get DOCX file size/metadata: %s", e, exc_info=True)
        
        return metadata
    
    def batch_extract(self, file_paths: List[str]) -> Dict[str, Dict[str, Any]]:
        """
        Extract text from multiple DOCX files.
        
        Args:
            file_paths: List of DOCX file paths
            
        Returns:
            Dictionary mapping file paths to extraction results
        """
        results = {}
        
        for file_path in file_paths:
            results[file_path] = self.extract_text(file_path)
        
        return results
    
    def is_supported(self, file_path: str) -> bool:
        """
        Check if file is supported by this extractor.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file is supported, False otherwise
        """
        return any(file_path.lower().endswith(ext) for ext in self.supported_extensions)
    
    def extract_structured_data(self, file_path: str) -> Dict[str, Any]:
        """
        Extract structured data from DOCX file (headings, lists, etc.).
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing structured data
        """
        if not DOCX_AVAILABLE:
            return {'error': 'DOCX library not available'}
        
        try:
            doc = Document(file_path)
            structured_data = {
                'headings': [],
                'paragraphs': [],
                'tables': [],
                'lists': []
            }
            
            # Extract headings (paragraphs with heading style)
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    structured_data['headings'].append({
                        'level': paragraph.style.name,
                        'text': paragraph.text.strip()
                    })
                elif paragraph.text.strip():
                    structured_data['paragraphs'].append(paragraph.text.strip())
            
            # Extract table data
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                structured_data['tables'].append(table_data)
            
            return structured_data
            
        except Exception as e:
            return {'error': f'Structured extraction failed: {str(e)}'}

    
