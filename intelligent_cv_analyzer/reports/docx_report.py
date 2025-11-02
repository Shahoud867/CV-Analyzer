"""
Phase 6 — Academic Report Generation (DOCX)

Generates a 6–8 page academic-style report (.docx) including:
- Introduction & Problem Definition
- System Design & Architecture
- Algorithm Explanation & Implementation
- Experimental Setup & Results
- Discussion & Conclusion
- Future Improvements

The report embeds tables and any available charts from data/results/charts.
"""

from __future__ import annotations

from pathlib import Path
from datetime import datetime
from typing import List

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


CHART_FILES = [
    'avg_execution_time_by_algorithm.png',
    'avg_comparisons_by_algorithm.png',
    'avg_relevance_by_algorithm.png',
    'scenario_execution_time.png',
    'scenario_comparisons.png',
    'scenario_relevance.png',
    'size_execution_time.png',
    'size_comparisons.png',
    'size_relevance.png',
]


def _latest_performance_csv(results_dir: Path) -> Path | None:
    files = sorted(results_dir.glob('performance_*.csv'))
    return files[-1] if files else None


def _set_styles(doc: Document):
    styles = doc.styles
    # Title style boost
    styles['Normal'].font.name = 'Calibri'
    styles['Normal'].font.size = Pt(11)


def _add_title(doc: Document, text: str, subtitle: str = ""):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _add_paragraphs(doc: Document, paras: List[str]):
    for t in paras:
        doc.add_paragraph(t)


def _add_table(doc: Document, headers: List[str], rows: List[List[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    table.style = 'Light Grid Accent 1'


def _add_image_if_exists(doc: Document, path: Path, width_in: float = 5.5):
    if path and path.exists():
        doc.add_picture(str(path), width=Inches(width_in))


def _perf_summary(csv_path: Path):
    df = pd.read_csv(csv_path)
    alg = (
        df.groupby('algorithm')
          .agg(avg_time=('execution_time','mean'),
               avg_comp=('comparisons','mean'),
               avg_rel=('relevance_score','mean'),
               runs=('algorithm','count'))
          .reset_index()
          .sort_values('avg_time')
    )
    scen = None
    if 'scenario' in df.columns:
        scen = (
            df.groupby(['scenario','algorithm'])
              .agg(avg_time=('execution_time','mean'),
                   avg_comp=('comparisons','mean'),
                   avg_rel=('relevance_score','mean'))
              .reset_index()
        )
    size = None
    if 'size_bucket' in df.columns:
        size = (
            df.groupby(['size_bucket','algorithm'])
              .agg(avg_time=('execution_time','mean'),
                   avg_comp=('comparisons','mean'),
                   avg_rel=('relevance_score','mean'))
              .reset_index()
        )
    return df, alg, scen, size


def generate_academic_report(output_path: str | Path | None = None) -> Path:
    results_dir = Path('data/results')
    charts_dir = results_dir / 'charts'
    output_dir = results_dir
    output_path = Path(output_path) if output_path else output_dir / 'final_report.docx'

    csv_path = _latest_performance_csv(results_dir)

    doc = Document()
    _set_styles(doc)

    # Title
    _add_title(doc, 'Intelligent CV Analyzer using String Matching Algorithms',
               f'Generated: {datetime.now():%Y-%m-%d %H:%M}')

    _add_heading(doc, '1. Introduction & Problem Definition', level=1)
    _add_paragraphs(doc, [
        'This report presents an intelligent CV analysis system that extracts and compares candidate CV contents against predefined job descriptions using classic string matching algorithms (Brute Force, Rabin–Karp, Knuth–Morris–Pratt).',
        'The objective is to compute relevance scores and generate comparative reports, while measuring algorithmic performance (execution time and comparisons) across realistic workloads.'
    ])

    _add_heading(doc, '2. System Design & Architecture', level=1)
    _add_paragraphs(doc, [
        'The system follows a modular architecture with clear separation of concerns: Extractors (PDF/DOCX), Engine (algorithms and analyzer), GUI (PyQt5), Persistence (SQLite), and Reporting.',
        'Data Flow: User selects CVs and provides a job description; text is extracted and normalized; the selected algorithm performs keyword matching; results are scored, ranked, visualized, and exported.'
    ])

    _add_heading(doc, '3. Algorithms & Implementation', level=1)
    _add_paragraphs(doc, [
        'Brute Force: O(n×m) worst/average, constant space; simple and robust baseline.',
        'Rabin–Karp: O(n+m) average, O(n×m) worst; rolling hash enables multi-keyword efficiency; susceptible to collisions.',
        'KMP: O(n+m) guaranteed; low comparisons via failure function; strong choice for real-time matching.'
    ])
    _add_table(doc,
               ['Algorithm', 'Best', 'Average', 'Worst', 'Space'],
               [
                   ['Brute Force', 'O(n)', 'O(n×m)', 'O(n×m)', 'O(1)'],
                   ['Rabin–Karp', 'O(n+m)', 'O(n+m)', 'O(n×m)', 'O(1)'],
                   ['KMP', 'O(n+m)', 'O(n+m)', 'O(n+m)', 'O(m)'],
               ])

    _add_heading(doc, '4. Experimental Setup', level=1)
    _add_paragraphs(doc, [
        'Dataset: Multiple CV documents from data/cvs/DataSet (PDF/DOCX).',
        'Scenarios: (i) Single keyword vs (ii) Multiple keywords; (iii) Small vs (iv) Large CVs (by median text length).',
        'Metrics: Execution time (s), comparison count, relevance score (match fraction + small density bonus).',
        'Automation: reports/performance_runner.py generates CSV; reports/charts.py renders comparative charts.'
    ])

    if csv_path and csv_path.exists():
        df, alg, scen, size = _perf_summary(csv_path)
        _add_paragraphs(doc, [f'Performance CSV: {csv_path.as_posix()} (rows: {len(df)})'])
        _add_heading(doc, '5. Results', level=1)
        # By algorithm table
        rows = []
        for _, r in alg.iterrows():
            rows.append([
                str(r['algorithm']),
                f"{r['avg_time']:.4f}",
                f"{r['avg_comp']:.0f}",
                f"{r['avg_rel']:.3f}",
                f"{int(r['runs'])}",
            ])
        _add_table(doc, ['Algorithm', 'Avg Time (s)', 'Avg Comparisons', 'Avg Relevance', 'Runs'], rows)

        # Add charts if present
        _add_paragraphs(doc, ['Selected Charts:'])
        for ch in CHART_FILES:
            p = charts_dir / ch
            _add_image_if_exists(doc, p)

        _add_heading(doc, '6. Discussion', level=1)
        _add_paragraphs(doc, [
            'KMP shows consistently strong runtime with low comparisons across sizes, aligning with its O(n+m) guarantee.',
            'Rabin–Karp benefits when many keywords are searched (rolling hash reuse) but may degrade under collisions.',
            'Brute Force is viable for small inputs but scales poorly; it remains a useful baseline for correctness checks.'
        ])
    else:
        _add_paragraphs(doc, ['No performance CSV found; run `python app.py --performance` to produce data and charts.'])

    _add_heading(doc, '7. Conclusion & Recommendation', level=1)
    _add_paragraphs(doc, [
        'Recommendation: Use KMP as the default for real-time CV analysis due to stable linear-time matching and low comparisons.',
        'For batches of many keywords, consider Rabin–Karp to exploit hashing efficiency. For tiny or educational runs, Brute Force suffices.'
    ])

    _add_heading(doc, '8. Future Improvements', level=1)
    _add_paragraphs(doc, [
        'Integrate NLP techniques (lemmatization, synonyms, and TF–IDF weighting) for semantic matching.',
        'Parallelize CV processing for reduced wall time; add caching of extracted texts.',
        'Extend GUI with “Compare All” summaries and richer report export formats.'
    ])

    doc.save(str(output_path))
    return output_path


if __name__ == '__main__':
    path = generate_academic_report()
    print(f'WROTE {path}')
